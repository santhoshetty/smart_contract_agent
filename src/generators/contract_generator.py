from pathlib import Path
from docx import Document
from datetime import date
from typing import Dict, Any
import logging

from ..validators.models import ContractData

logger = logging.getLogger(__name__)

class ContractGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path
        logger.info(f"Initialized ContractGenerator with template: {template_path}")

    def generate_contract(self, data: ContractData, output_path: str) -> None:
        """Generate a contract by replacing placeholders in the template with actual values."""
        try:
            logger.info(f"Loading template from: {self.template_path}")
            doc = Document(self.template_path)
            
            # Replace placeholders with actual values
            self._replace_placeholders(doc, data)
            
            # Save the generated document
            logger.info(f"Saving generated contract to: {output_path}")
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error generating contract: {str(e)}")
            raise

    def _replace_placeholders(self, doc: Document, data: ContractData):
        """Replace placeholders in the document with actual values."""
        # Convert data to dictionary for easier access
        data_dict = data.model_dump()
        logger.info(f"Replacing placeholders with data: {data_dict}")
        
        def replace_in_text(text: str, key: str, value: Any) -> str:
            """Helper function to replace all placeholder formats in text."""
            if isinstance(value, list):
                value = ", ".join(value)
            elif isinstance(value, date):
                value = value.strftime("%d %b, %Y")
            
            # Try all possible placeholder formats
            placeholders = [
                f"{{{key}}}",  # Single braces
                f"{{{{{key}}}}}",  # Double braces
                f"{{{key}}}",  # Single braces with spaces
                f"{{{{{key}}}}}",  # Double braces with spaces
                f"{{{{{key}}}}}",  # Double braces with spaces
                f"{{{key}}}",  # Single braces with spaces
            ]
            
            for placeholder in placeholders:
                if placeholder in text:
                    logger.info(f"Replacing {placeholder} with {value} in text")
                    # Convert value to string and ensure no braces
                    value_str = str(value).strip('{}')
                    # Replace the placeholder with the value
                    text = text.replace(placeholder, value_str)
                    # Remove any extra braces that might have been added
                    text = text.replace(f"{{{value_str}}}", value_str)
            return text
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for key, value in data_dict.items():
                text = replace_in_text(text, key, value)
            paragraph.text = text
        
        # Replace placeholders in tables
        for table in doc.tables:
            # First, check if this is the machine table
            is_machine_table = False
            for row in table.rows:
                for cell in row.cells:
                    if any("machine" in p.text.lower() for p in cell.paragraphs):
                        is_machine_table = True
                        break
                if is_machine_table:
                    break
            
            if is_machine_table and data_dict.get('machine_names'):
                # Handle machine table specially
                machine_names = data_dict['machine_names']
                logger.info(f"Processing machine table with names: {machine_names}")
                # Get the template row (first data row)
                template_row = table.rows[1] if len(table.rows) > 1 else None
                
                if template_row:
                    # Remove existing rows except header
                    for _ in range(len(table.rows) - 1):
                        table._element.remove(table.rows[-1]._element)
                    
                    # Add a new row for each machine
                    for machine in machine_names:
                        new_row = table.add_row()
                        for i, cell in enumerate(template_row.cells):
                            # Copy the template cell content
                            new_cell = new_row.cells[i]
                            for p in cell.paragraphs:
                                # Replace all placeholders in the cell
                                text = p.text
                                # First replace machine_name
                                text = replace_in_text(text, "machine_name", machine)
                                # Then replace other placeholders
                                for key, value in data_dict.items():
                                    text = replace_in_text(text, key, value)
                                logger.info(f"Setting cell text to: {text}")
                                new_cell.text = text
            else:
                # Handle regular table cells
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            for key, value in data_dict.items():
                                text = replace_in_text(text, key, value)
                            paragraph.text = text 